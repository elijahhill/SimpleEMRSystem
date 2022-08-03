import json
import re

from pathlib import Path
from docx import Document

from os import listdir as os_listdir
from os.path import isdir as os_isdir


class Notes:

    def __init__(self, progress_notes_path):
        self.patient_folder_paths = {}
        self.id_matcher = re.compile("\d{4,}")
        self.__init_patient_folder_paths__(progress_notes_path=progress_notes_path)
        self.history_and_physical = re.compile(
            "((h|H)\s*(&)\s*(p|P))|(History\s(&|And)\sPhysical)", flags=re.IGNORECASE)
        self.day_matcher = re.compile("(day)(\s*)(\d)", re.IGNORECASE)

    def __init_patient_folder_paths__(self, progress_notes_path: str):
        overall_note_folder = os_listdir(progress_notes_path)
        for case_note_folder in overall_note_folder:
            potential_case_folder = Path(
                progress_notes_path) / case_note_folder
            if os_isdir(potential_case_folder):
                patient_id_match = self.id_matcher.search(
                    potential_case_folder.stem)
                if patient_id_match != None:
                    patient_id = int(patient_id_match.group(0))
                    self.patient_folder_paths[patient_id] = potential_case_folder

    def create_notes(self, data_layout, patient_path, case_id):
        all_notes_dict = {}

        note_date = "11/11"
        note_text = "DEFAULT TEXT"
        note_time = 1352687800000.0
        note_headers = data_layout["notes"]

        if case_id in self.patient_folder_paths:

            doc_files = []
            
            # Filter out document files
            folder_path = self.patient_folder_paths[case_id]
            for potential_note in os_listdir(folder_path):
                if potential_note.endswith(".docx"):
                    doc_files.append(potential_note)

            h_and_p = ''
            entry_two = ''
            entry_three = ''

            # Assign searched items to their particular slots
            added = set()
            for doc_file in doc_files:
                h_and_p_search = self.history_and_physical.search(doc_file)
                if h_and_p_search != None:
                    h_and_p = h_and_p_search
                    added.add(h_and_p.string)
                day_search = self.day_matcher.search(doc_file)

                day_search_exists = day_search != None
                if day_search_exists:
                    already_added = day_search.string in added
                    if day_search_exists and not already_added:
                        for index, entry in enumerate([h_and_p, entry_two, entry_three]):
                            # Place the new search entry in the first available blank variable
                            if entry == '':
                                if index == 0:
                                    h_and_p = day_search
                                    added.add(day_search.string)
                                elif index == 1:
                                    entry_two = day_search
                                    added.add(day_search.string)
                                elif index == 2:
                                    entry_three = day_search
                                    added.add(day_search.string)
                                break



            if h_and_p != '':
                h_and_p_type = "H&P"
                note_dict = {
                    "date": note_date,
                    "text": note_text,
                    "js_time": note_time,
                    "upk": 0,
                    "type": h_and_p_type
                }

                with open(f'{self.patient_folder_paths[case_id]}/{h_and_p.string}', 'rb') as fp:
                    document = Document(fp)
                    note_dict['text'] = '\n'.join([paragraph.text for paragraph in document.paragraphs])

                all_notes_dict[h_and_p_type] = note_dict

            if entry_two != '':
                second_entry_type = f'Day{entry_two.group(len(entry_two.groups()))}'
                note_dict = {
                    "date": note_date,
                    "text": note_text,
                    "js_time": note_time,
                    "upk": 0,
                    "type": second_entry_type
                }
                with open(f'{self.patient_folder_paths[case_id]}/{entry_two.string}', 'rb') as fp:
                    document = Document(fp)
                    note_dict['text'] = '\n'.join(
                        [paragraph.text for paragraph in document.paragraphs])

                all_notes_dict[second_entry_type] = note_dict
            
            if entry_three != '':
                third_entry_type = f'Day{entry_three.group(len(entry_three.groups()))}'
                note_dict = {
                    "date": note_date,
                    "text": note_text,
                    "js_time": note_time,
                    "upk": 0,
                    "type": third_entry_type
                }

                with open(f'{self.patient_folder_paths[case_id]}/{entry_three.string}', 'rb') as fp:
                    document = Document(fp)
                    note_dict['text'] = '\n'.join(
                        [paragraph.text for paragraph in document.paragraphs])

                all_notes_dict[third_entry_type] = note_dict

        try:
            note_panel_path = patient_path / "note_panel_data.json"
            note_panel_path.touch(exist_ok=True)
        except PermissionError:
            raise PermissionError(
                "Permission denied when creating notes file")

        with open(note_panel_path, "w+") as fp:
            json.dump(obj=all_notes_dict, fp=fp, indent=4)
